from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from docx.table import Table


BASE = Path(r"C:\Users\OWNER\Desktop\GPP")
DOCX_PATH = BASE / "GPP 보고서.docx"
OUT_PATH = BASE / "GPP 보고서_HalloweenParty작성.docx"
WORKSPACE_OUT = Path(r"C:\Users\OWNER\Documents\GitHub\Personal_Study\GameEnginePractice\GPP 보고서_HalloweenParty작성.docx")

IMAGES = {
    "before_script": BASE / "할로윈파티" / "수정전" / "Script가 가장 높을때.png",
    "before_gc": BASE / "할로윈파티" / "수정전" / "GC가 가장 높을때.png",
    "before_memory": BASE / "할로윈파티" / "수정전" / "메모리가 가장 높을떄.png",
    "after_script": BASE / "할로윈파티" / "수정후" / "SCript가 가장 높을 떄.png",
    "after_gc": BASE / "할로윈파티" / "수정후" / "GC가 가장 높을 때.png",
    "after_memory": BASE / "할로윈파티" / "수정후" / "메모리가 가장 높을 때.png",
}


def iter_blocks(doc):
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()


def set_cell_text(cell, text, bold=False):
    clear_cell(cell)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_image_cell(cell, label, image_path):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(8)
    p.add_run("\n")
    if image_path.exists():
        p.add_run().add_picture(str(image_path), width=Inches(2.25))
    else:
        p.add_run(f"[이미지 없음: {image_path.name}]")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def insert_paragraph_after(block, text, style=None, bold_prefix=None):
    new_p = OxmlElement("w:p")
    block._element.addnext(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if style:
        paragraph.style = style
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    return paragraph


def find_paragraph(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"paragraph not found: {text}")


def find_next_table_after(doc, paragraph_text):
    blocks = list(iter_blocks(doc))
    for i, block in enumerate(blocks):
        if isinstance(block, Paragraph) and block.text.strip() == paragraph_text:
            for next_block in blocks[i + 1:]:
                if isinstance(next_block, Table):
                    return next_block
    raise ValueError(f"table after paragraph not found: {paragraph_text}")


def fill_profiler_image_table(table):
    rows = [
        ("Script가 가장 높을 때", "수정전 / Script가 가장 높을때.png", IMAGES["before_script"], "수정후 / SCript가 가장 높을 떄.png", IMAGES["after_script"]),
        ("GC가 가장 높을 때", "수정전 / GC가 가장 높을때.png", IMAGES["before_gc"], "수정후 / GC가 가장 높을 때.png", IMAGES["after_gc"]),
        ("메모리가 가장 높을 때", "수정전 / 메모리가 가장 높을떄.png", IMAGES["before_memory"], "수정후 / 메모리가 가장 높을 때.png", IMAGES["after_memory"]),
    ]

    for cell, text in zip(table.rows[0].cells, ["유형", "수정 전", "수정 후"]):
        set_cell_text(cell, text, bold=True)

    for row_idx, row in enumerate(rows, start=1):
        kind, before_label, before_image, after_label, after_image = row
        set_cell_text(table.cell(row_idx, 0), kind, bold=True)
        add_image_cell(table.cell(row_idx, 1), before_label, before_image)
        add_image_cell(table.cell(row_idx, 2), after_label, after_image)


def fill_comparison_table(table):
    data = [
        ("CPU Frame Time", "약 6.87~6.95ms", "약 6.88~6.92ms"),
        ("Script 최고 프레임", "Animator/Canvas 처리 중 GC.Alloc 344~414B 발생", "Script 구간 GC Alloc 0B, ScriptRunBehaviourUpdate 약 0.07ms"),
        ("GC Alloc 최고 프레임", "13.8KB, NPCManager.MakeNPC -> Instantiate 약 12.0KB", "0B, GarbageCollector.CollectIncremental 약 0.02ms"),
        ("오브젝트 생성 방식", "NPC/Enemy 생성 시 Instantiate, 일부 Destroy/System.Collections 호출", "Pool에서 미리 생성한 오브젝트를 재사용"),
        ("Total Committed Memory", "336.4MB (Max 336.5MB)", "272.2MB (Max 272.3MB)"),
        ("Managed Heap", "1.8 / 2.6MB", "1.8 / 2.0MB"),
        ("Object Count", "약 5.49k~5.79k", "약 8.34k~8.88k"),
        ("Rendering", "Batches 464~642, SetPass 124~221", "Batches 398~417, SetPass 57~69"),
        ("Triangles / Vertices", "약 133.26k~232.80k / 167.16k~290.09k", "약 131.85k~155.35k / 163.86k~179.81k"),
    ]

    for cell, text in zip(table.rows[0].cells, ["측정 항목", "수정 전", "수정 후"]):
        set_cell_text(cell, text, bold=True)

    for row_idx, row in enumerate(data, start=1):
        for col_idx, text in enumerate(row):
            set_cell_text(table.cell(row_idx, col_idx), text, bold=(col_idx == 0))


def add_analysis_sections(doc, profiler_table, comparison_table):
    p25_text = (
        "Profiler 측정은 수정 전과 수정 후를 각각 Script, GC Alloc, Memory가 가장 높게 나타난 프레임으로 나누어 캡처하였다. "
        "수정 전에는 NPCManager.MakeNPC 코루틴에서 Instantiate가 호출되며 GC Alloc이 크게 발생했고, Animator/Canvas 처리 중에도 "
        "소량의 할당이 반복되었다. 수정 후에는 NPC, Special NPC, Enemy를 게임 시작 시 Pool에 미리 생성하고 재사용하도록 변경하여 "
        "실행 중 Instantiate로 인한 할당이 사라졌다."
    )
    insert_paragraph_after(profiler_table, p25_text)

    p26_1 = (
        "수정 전 GC가 가장 높은 프레임에서는 PlayerLoop 기준 13.8KB의 GC Alloc이 발생했으며, 그중 NPCManager.MakeNPC 코루틴의 "
        "Instantiate 호출이 약 12.0KB를 차지하였다. 이는 NPC가 새로 등장할 때마다 런타임 객체 생성 비용과 메모리 할당이 함께 발생했다는 "
        "의미이다. 수정 후 GC 최고 프레임에서는 GC Alloc이 0B로 측정되어, 게임 진행 중 반복 생성으로 인한 할당이 제거되었음을 확인할 수 있다."
    )
    p26_2 = (
        "CPU Frame Time은 수정 전후 모두 약 6.9ms 수준으로 큰 차이는 없었다. 그러나 수정 후 Script 구간의 GC Alloc이 0B로 유지되고, "
        "Batches와 SetPass Calls가 전반적으로 감소하여 프레임 안정성이 좋아졌다. Object Count는 수정 후 증가했는데, 이는 Pool에 사용할 "
        "NPC와 Enemy를 미리 생성해 비활성 상태로 보관하기 때문이다. 즉, 메모리에 미리 올려두는 대신 플레이 중 생성/삭제 비용과 GC 스파이크를 "
        "줄이는 방향으로 개선되었다."
    )
    insert_paragraph_after(comparison_table, p26_2)
    insert_paragraph_after(comparison_table, p26_1)

    heading_27 = find_paragraph(doc, "2.7 개선 결과 요약")
    summary = [
        "Halloween Party의 주요 병목은 NPC와 Enemy를 필요할 때마다 Instantiate/Destroy하는 구조에서 발생하였다. 개선 후에는 NPCPool과 EnemyPool을 통해 게임 시작 시 필요한 객체를 미리 생성하고, 이후에는 SetActive 기반으로 재사용하도록 변경하였다.",
        "그 결과 수정 전 GC 최고 프레임에서 13.8KB까지 발생하던 할당이 수정 후 0B로 감소하였다. 특히 NPCManager.MakeNPC에서 발생하던 Instantiate 12.0KB 할당이 제거되어, NPC 등장 시점의 GC 스파이크가 사라졌다.",
        "Memory 측면에서는 Object Count가 증가했지만 이는 Pool에 의해 의도적으로 보관되는 객체 수가 늘어난 결과이다. 대신 Total Committed Memory는 336.4MB에서 272.2MB로 감소했고, 실행 중 할당이 줄어 전체적인 프레임 안정성이 향상되었다.",
        "따라서 이번 개선은 평균 CPU 시간을 크게 낮추는 것보다, 플레이 중 발생하는 순간적인 객체 생성 비용과 GC Alloc을 줄여 게임 진행 중 끊김 가능성을 낮추는 데 효과가 있었다.",
    ]
    anchor = heading_27
    for text in reversed(summary):
        insert_paragraph_after(anchor, text)


def main():
    doc = Document(str(DOCX_PATH))
    profiler_table = find_next_table_after(doc, "2.5 Profiler 측정 결과,")
    comparison_table = find_next_table_after(doc, "2.6 수정 전후 성능 비교 분석")

    fill_profiler_image_table(profiler_table)
    fill_comparison_table(comparison_table)
    add_analysis_sections(doc, profiler_table, comparison_table)

    doc.save(str(WORKSPACE_OUT))
    print(WORKSPACE_OUT)


if __name__ == "__main__":
    main()
